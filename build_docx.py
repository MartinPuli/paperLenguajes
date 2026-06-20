#!/usr/bin/env python3
"""Convierte paper.md en un .docx con formato académico (portada, cuerpo justificado,
bibliografía con sangría francesa). Uso: python3 build_docx.py"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "paper.md"
OUT = "Paper_Rol_del_ingeniero_informatico_en_10_anios.docx"

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*)')
IMG = re.compile(r'^!\[(.*)\]\(([^()]+)\)\s*$')

def add_runs(paragraph, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        seg = m.group(0)
        if seg.startswith('**'):
            r = paragraph.add_run(seg[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(seg[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

# --- read & split markdown -------------------------------------------------
lines = open(SRC, encoding="utf-8").read().split("\n")
title = lines[0].lstrip("# ").strip()
body = lines[1:]

doc = Document()

# base style: serif, 12pt
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)

# --- title page ------------------------------------------------------------
def centered(text, *, size=12, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p

centered("UNIVERSIDAD DEL CEMA", size=14, bold=True, before=48)
centered("Buenos Aires, Argentina", size=12, after=120)
centered(title, size=16, bold=True, before=60, after=60)
centered("Trabajo Práctico Final", size=12, after=160)
centered("Carrera: Ingeniería Informática (3.er año)", size=12, before=40)
centered("Materia: Lenguajes Formales y Teoría de la Computación", size=12)
centered("Profesor/a: ______________________________", size=12)
centered("Integrantes: Martín Ezequiel Pulitano y Nicolás Silva", size=12)
centered("Junio de 2026", size=12, before=120)
doc.add_page_break()

# --- body ------------------------------------------------------------------
in_biblio = False
for raw in body:
    line = raw.rstrip()
    if not line.strip():
        continue
    if line.startswith("## "):
        text = line[3:].strip()
        in_biblio = text.lower().startswith(("bibliograf", "referencia"))
        h = doc.add_heading(level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = h.add_run(text)
        r.font.name = "Times New Roman"; r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 0, 0); r.bold = True
        continue
    if line.startswith("### "):
        h = doc.add_heading(level=2)
        r = h.add_run(line[4:].strip())
        r.font.name = "Times New Roman"; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0); r.bold = True
        continue
    mimg = IMG.match(line)
    if mimg:
        caption, path = mimg.group(1), mimg.group(2)
        doc.add_picture(path, width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_before = Pt(8)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        add_runs(cap, caption)
        for r in cap.runs:
            r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        continue
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    if in_biblio:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)  # hanging indent
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0.4)
    add_runs(p, line.strip())

# --- footer con número de página (excepto en la portada) -------------------
section = doc.sections[0]
section.different_first_page_header_footer = True
fp = section.first_page_footer
fp.paragraphs[0].text = ""
footer = section.footer
fpar = footer.paragraphs[0]
fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fpar.add_run()
b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
run._r.append(b); run._r.append(instr); run._r.append(e)

doc.save(OUT)
print("Escrito:", OUT)
